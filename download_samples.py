import os
import requests
from pathlib import Path
import json
import time

def download_file(url: str, output_path: str):
    """Download a file from URL to the specified path."""
    response = requests.get(url, stream=True)
    response.raise_for_status()
    
    with open(output_path, 'wb') as f:
        for chunk in response.iter_content(chunk_size=8192):
            f.write(chunk)

def create_sample_documents():
    """Create sample documents directory and download example files."""
    # Create directories
    sample_dir = Path("sample_documents")
    sample_dir.mkdir(exist_ok=True)
    
    # Sample documents to download
    documents = [
        {
            "name": "contract.pdf",
            "url": "https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf",
            "type": "legal"
        },
        {
            "name": "technical_doc.pdf",
            "url": "https://www.ietf.org/rfc/rfc2616.txt",
            "type": "technical"
        },
        {
            "name": "financial_report.pdf",
            "url": "https://www.adobe.com/content/dam/acom/en/documents/pdf/accessibility/accessibility-pdf-ua.pdf",
            "type": "financial"
        }
    ]
    
    print("Downloading sample documents...")
    
    for doc in documents:
        output_path = sample_dir / doc["name"]
        if not output_path.exists():
            try:
                print(f"Downloading {doc['name']}...")
                download_file(doc["url"], output_path)
                print(f"Successfully downloaded {doc['name']}")
                time.sleep(1)  # Be nice to the servers
            except Exception as e:
                print(f"Error downloading {doc['name']}: {str(e)}")
        else:
            print(f"{doc['name']} already exists")
    
    # Create a sample HTML document
    html_content = """
    <html>
    <head>
        <title>Sample HTML Document</title>
    </head>
    <body>
        <h1>Sample HTML Document</h1>
        <p>This is a sample HTML document for testing the document processing pipeline.</p>
        <h2>Section 1</h2>
        <p>This section contains some sample text for testing metadata extraction.</p>
        <h2>Section 2</h2>
        <p>This section contains more sample text for testing document classification.</p>
    </body>
    </html>
    """
    
    with open(sample_dir / "sample.html", "w") as f:
        f.write(html_content)
    
    # Create a sample DOCX document
    from docx import Document
    doc = Document()
    doc.add_heading('Sample DOCX Document', 0)
    doc.add_paragraph('This is a sample DOCX document for testing the document processing pipeline.')
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('This section contains some sample text for testing metadata extraction.')
    doc.add_heading('Section 2', level=1)
    doc.add_paragraph('This section contains more sample text for testing document classification.')
    doc.save(sample_dir / "sample.docx")
    
    print("\nSample documents have been created in the 'sample_documents' directory:")
    for file in sample_dir.iterdir():
        print(f"- {file.name}")

if __name__ == "__main__":
    create_sample_documents() 